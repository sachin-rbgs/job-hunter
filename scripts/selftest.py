"""Offline self-test of the parts that do not need network or a database.

Checks CV loading, format-preserving .docx edits, ATS alignment maths, UK red-flag
detection, and the anti-fabrication guard. Run this after changing scoring or editor
logic.

  python scripts/selftest.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from app import config, verify  # noqa: E402
from app.cv import editor, loader  # noqa: E402
from app.scoring import ats_alignment, detect_red_flags, extract_required_years, score_job  # noqa: E402

CV_ROOT = Path(__file__).resolve().parent.parent.parent
passed = failed = 0


def check(name, condition, detail=""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  PASS  {name}")
    else:
        failed += 1
        print(f"  FAIL  {name} {detail}")


print("\n1. CV variant discovery")
variants = loader.discover(CV_ROOT)
check("found variants", len(variants) >= 3, f"got {len(variants)}")
for v in variants:
    print(f"        {v.name:18} {len(v.bullets):3} bullets, {len(v.full_text):5} chars")
check("no India CV", all("india" not in v.name.lower() for v in variants))
check("all have bullets", all(len(v.bullets) > 0 for v in variants))

master = loader.load_master_profile(CV_ROOT)
check("master profile loaded", len(master) > 1000, f"{len(master)} chars")


print("\n2. Format-preserving .docx edit")
if variants:
    src = Path(variants[0].file_path)
    original_doc = Document(str(src))
    target = variants[0].bullets[0]
    idx, old_text = target["index"], target["text"]

    before_para = original_doc.paragraphs[idx]
    before_style = before_para.style.name
    before_font = before_para.runs[0].font.name if before_para.runs else None
    before_size = before_para.runs[0].font.size if before_para.runs else None
    before_bold = before_para.runs[0].font.bold if before_para.runs else None
    before_count = len(original_doc.paragraphs)

    # Same length so the strict length guard accepts it.
    new_text = old_text[:-10] + "X" * 10
    out = config.OUTPUT_DIR / "_selftest.docx"
    report = editor.apply_edits(src, out, [{"index": idx, "original": old_text, "revised": new_text}])
    check("edit applied", len(report["applied"]) == 1, str(report.get("rejected")))

    after_doc = Document(str(out))
    after_para = after_doc.paragraphs[idx]
    check("paragraph count unchanged", len(after_doc.paragraphs) == before_count)
    check("style preserved", after_para.style.name == before_style)
    check("font preserved", (after_para.runs[0].font.name if after_para.runs else None) == before_font)
    check("size preserved", (after_para.runs[0].font.size if after_para.runs else None) == before_size)
    check("bold preserved", (after_para.runs[0].font.bold if after_para.runs else None) == before_bold)
    check("text changed", after_para.text.strip().endswith("XXXXXXXXXX"))

    # Length guard must reject an over-long rewrite.
    rejected = editor.apply_edits(
        src, config.OUTPUT_DIR / "_selftest2.docx",
        [{"index": idx, "original": old_text, "revised": old_text + " " + "padding " * 30}],
    )
    check("length guard rejects drift", len(rejected["rejected"]) == 1)

    # Wrong original text must be refused, guarding against index drift.
    drifted = editor.apply_edits(
        src, config.OUTPUT_DIR / "_selftest3.docx",
        [{"index": idx, "original": "completely different text here", "revised": "nope"}],
    )
    check("drift guard rejects mismatch", len(drifted["rejected"]) == 1)


print("\n3. ATS alignment")
jd = """We seek a graduate mechanical engineer with SolidWorks and ANSYS experience.
Knowledge of GD&T, DFMEA and tolerance stack up is essential. Familiarity with
Minitab, Six Sigma and IATF 16949 preferred. CATIA experience a bonus."""
if variants:
    result = ats_alignment(jd, variants[0].full_text)
    print(f"        required: {result['required']}")
    print(f"        matched:  {result['matched']}")
    print(f"        missing:  {result['missing']}")
    print(f"        score:    {result['score']}%")
    check("alignment computed", 0 <= result["score"] <= 100)
    check("finds real terms", "solidworks" in result["matched"] or "solidworks" in result["missing"])

empty = ats_alignment("", "anything")
check("no requirements scores 100", empty["score"] == 100)


print("\n4. Honest keyword lift")
honest, rejected = verify.verify_skill_additions(["catia", "minitab", "six sigma"], master)
print(f"        honest:   {honest}")
print(f"        rejected: {rejected}")
check("rejects CATIA (not in profile)", "catia" not in [h.lower() for h in honest])
check("accepts something real", len(honest) >= 1)


print("\n5. UK red flags")
cases = [
    ({"title": "Junior Design Engineer", "company": "X",
      "description": "You will need a minimum of 4 years experience in design."},
     "junior_title_but_4y_required"),
    ({"title": "Mechanical Engineer", "company": "X",
      "description": "Candidates must hold active SC clearance."},
     "security_clearance_required"),
    ({"title": "Senior Design Engineer", "company": "X", "description": "Lead the team."},
     "seniority_mismatch"),
    ({"title": "Design Engineer", "company": "Blue Arrow Recruitment", "description": "Great role."},
     "agency_listing"),
    ({"title": "Design Engineer", "company": "X",
      "description": "Applicants must be a Chartered Engineer (CEng)."},
     "chartership_required"),
]
for job, expected in cases:
    flags = detect_red_flags(job)
    check(f"detects {expected}", expected in flags, f"got {flags}")

check("parses '4 years'", extract_required_years("minimum of 4 years experience") == 4)
check("parses '5+ years'", extract_required_years("5+ years experience required") == 5)
check("ignores absent", extract_required_years("no requirement stated") is None)


print("\n6. Scoring end to end")
if variants:
    corpus = "\n".join(v.full_text for v in variants)
    good = score_job({
        "title": "Graduate Mechanical Design Engineer",
        "company": "Rolls-Royce", "location": "Derby",
        "commute_miles": 18, "salary_min": 32000,
        "description": jd,
    }, corpus)
    bad = score_job({
        "title": "Senior Principal Engineer",
        "company": "Acme Recruitment", "location": "London",
        "commute_miles": 110, "salary_min": None,
        "description": "10+ years experience required. SC clearance essential.",
    }, corpus)
    print(f"        graduate role: {good['score']}  {good['breakdown']}")
    print(f"        senior role:   {bad['score']}  flags={bad['red_flags']}")
    check("graduate role scores well", good["score"] >= 55, str(good["score"]))
    check("clearance role zeroed", bad["score"] == 0, str(bad["score"]))


print("\n7. Fabrication guard")
sources = [variants[0].full_text if variants else "", master]
clean = verify.check("Applied SolidWorks and Minitab to validation work.", sources)
dirty = verify.check(
    "I am a Chartered Engineer (CEng) who worked at Fictional Dynamics Ltd in 1994, "
    "improving throughput by 87%.", sources)
print(f"        clean issues: {clean}")
print(f"        dirty issues: {dirty}")
check("clean text passes", len(clean) == 0, str(clean))
check("catches fake credential", any("ceng" in i.lower() for i in dirty))
check("catches fake year", any("1994" in i for i in dirty))
check("catches fake company", any("Fictional Dynamics" in i for i in dirty))
check("catches fake metric", any("87" in i for i in dirty))


print(f"\n{'='*46}\n  {passed} passed, {failed} failed\n{'='*46}")
for tmp in config.OUTPUT_DIR.glob("_selftest*.docx"):
    try:
        tmp.unlink(missing_ok=True)
    except OSError:
        pass
sys.exit(1 if failed else 0)
