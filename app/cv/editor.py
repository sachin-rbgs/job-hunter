"""Format-preserving .docx editing.

The rule: never delete and re-add a paragraph. Rewrite the text inside its existing
runs. Runs carry font, size, bold, colour and spacing, so editing `run.text` keeps the
document visually identical apart from the words themselves.
"""
from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

LENGTH_TOLERANCE = 0.15  # keep rewrites within +/-15% so line wrapping does not shift


class LengthDriftError(ValueError):
    pass


def check_length(original: str, revised: str, tolerance: float = LENGTH_TOLERANCE) -> bool:
    if not original:
        return True
    ratio = abs(len(revised) - len(original)) / len(original)
    return ratio <= tolerance


def set_paragraph_text(para: Paragraph, new_text: str) -> None:
    """Replace a paragraph's text, preserving the formatting of its first run.

    Any leading bullet glyph or whitespace in the original is retained so list
    indentation and manual bullets survive.
    """
    if not para.runs:
        para.add_run(new_text)
        return

    prefix_match = re.match(r"^\s*[•●\-\*▪–]?\s*", para.text)
    prefix = prefix_match.group(0) if prefix_match else ""

    first = para.runs[0]
    first.text = prefix + new_text.strip()
    for run in para.runs[1:]:
        run.text = ""


def append_to_paragraph(para: Paragraph, addition: str) -> None:
    """Append text to a paragraph by extending its last non-empty run.

    Used to add a true skill to an existing skills line. Extending an existing run
    means the new text inherits that run's formatting exactly.
    """
    for run in reversed(para.runs):
        if run.text.strip():
            run.text = run.text.rstrip() + addition
            return
    para.add_run(addition)


def duplicate_paragraph_after(para: Paragraph, new_text: str) -> Paragraph:
    """Clone a paragraph (keeping style and run formatting) and set new text.

    Cloning an existing bullet is the only safe way to add one: it inherits the
    list style, indentation and numbering context of its neighbour.
    """
    new_element = copy.deepcopy(para._p)
    para._p.addnext(new_element)
    new_para = Paragraph(new_element, para._parent)
    set_paragraph_text(new_para, new_text)
    return new_para


def apply_edits(
    source_path: str | Path,
    output_path: str | Path,
    bullet_edits: list[dict],
    skill_additions: list[dict] | None = None,
    strict_length: bool = True,
) -> dict:
    """Write an edited copy of a CV.

    bullet_edits:    [{"index": int, "original": str, "revised": str}]
    skill_additions: [{"index": int, "addition": ", ANSYS Workbench"}]

    Returns a report of what was applied and what was rejected.
    """
    doc = Document(str(source_path))
    paragraphs = doc.paragraphs
    report = {"applied": [], "rejected": []}

    for edit in bullet_edits or []:
        idx = edit.get("index")
        revised = (edit.get("revised") or "").strip()
        original = edit.get("original") or ""
        if idx is None or not revised or idx >= len(paragraphs):
            report["rejected"].append({**edit, "reason": "bad index or empty text"})
            continue

        para = paragraphs[idx]
        # Guard against editing the wrong paragraph if the doc changed underneath us.
        if original and para.text.strip()[:40] != original.strip()[:40]:
            report["rejected"].append({**edit, "reason": "paragraph text drifted"})
            continue
        if strict_length and not check_length(para.text.strip(), revised):
            report["rejected"].append({**edit, "reason": "length drift over 15%"})
            continue

        set_paragraph_text(para, revised)
        report["applied"].append({"index": idx, "revised": revised})

    for addition in skill_additions or []:
        idx = addition.get("index")
        text = addition.get("addition", "")
        if idx is None or not text or idx >= len(paragraphs):
            report["rejected"].append({**addition, "reason": "bad index"})
            continue
        append_to_paragraph(paragraphs[idx], text)
        report["applied"].append({"index": idx, "added": text})

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    report["output"] = str(output_path)
    return report


def find_skills_paragraphs(source_path: str | Path) -> list[dict]:
    """Locate paragraphs that read like skill lists, as targets for honest additions."""
    doc = Document(str(source_path))
    hits = []
    markers = ("skill", "software", "tools", "technical", "competenc", "proficien")
    for i, para in enumerate(doc.paragraphs):
        low = para.text.lower()
        if not low.strip():
            continue
        if any(m in low for m in markers) or low.count(",") >= 3:
            hits.append({"index": i, "text": para.text.strip()})
    return hits
