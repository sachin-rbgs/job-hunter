"""Read .docx CV variants into structured form."""
from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from app import config

# Folder name in the CV directory -> (variant name, lean tags)
VARIANT_MAP = {
    ".": ("General", ["general"]),
    "FEA CV": ("FEA", ["simulation", "fea", "cae", "thermal", "structural"]),
    "Process Quality based CV": ("Process Quality", ["quality", "iso", "iatf", "spc", "rcca", "audit"]),
    "Design Engineer CV": ("Design Engineer", ["cad", "solidworks", "gdt", "drawings", "dfm"]),
    "NPI&R&D": ("NPI & R&D", ["npi", "dfmea", "pfmea", "product development", "r&d"]),
    "Commercial Client Facing": ("Commercial", ["customer-facing", "application engineering", "commercial"]),
}

# Never routed to.
EXCLUDED = {"India CV"}

BULLET_RE = re.compile(r"^\s*[•●\-\*▪–]\s*")


@dataclass
class LoadedCV:
    name: str
    file_path: str
    full_text: str
    bullets: list = field(default_factory=list)
    lean_tags: list = field(default_factory=list)


def _is_bullet(para) -> bool:
    """A paragraph counts as an editable bullet if it is list-styled or long prose."""
    style = (para.style.name or "").lower()
    text = para.text.strip()
    if not text or len(text) < 40:
        return False
    if "list" in style or "bullet" in style:
        return True
    if BULLET_RE.match(para.text):
        return True
    # Experience bullets in these CVs are often plain paragraphs starting with a verb.
    return len(text) > 60 and text[0].isupper() and not text.endswith(":")


def extract(path: Path) -> tuple[str, list]:
    """Return (full_text, bullets) for a .docx. Bullet index maps to doc.paragraphs."""
    doc = Document(str(path))
    lines, bullets = [], []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            lines.append(text)
        if _is_bullet(para):
            bullets.append({"index": i, "text": text})
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines), bullets


def discover(cv_root: Path, verbose: bool = True) -> list[LoadedCV]:
    """Find CV variants under the user's CV folder and copy them into data/cv_variants.

    The CV folder is OneDrive-synced, so some files may be cloud-only placeholders that
    raise OSError on read. Those are skipped with a message rather than killing the run:
    open the file once in Explorer, or right-click the folder and choose
    "Always keep on this device", then re-run.
    """
    found: list[LoadedCV] = []
    for folder, (name, tags) in VARIANT_MAP.items():
        directory = cv_root if folder == "." else cv_root / folder
        if not directory.is_dir():
            continue

        candidates = [c for c in sorted(directory.glob("*.docx"))
                      if not c.name.startswith("~$")]
        if not candidates:
            if verbose:
                print(f"  skip {name}: no .docx in {directory.name} "
                      f"(a PDF-only folder cannot be edited)")
            continue

        # Prefer a file with "CV" in the name over cover letters.
        cvs = [c for c in candidates if "cv" in c.name.lower()]
        src = (cvs or candidates)[0]
        dest = config.CV_DIR / f"{name.replace(' ', '_').replace('&', 'and')}.docx"

        try:
            shutil.copy2(src, dest)
            full_text, bullets = extract(dest)
        except OSError as exc:
            if verbose:
                print(f"  skip {name}: cannot read {src.name} ({exc.strerror or exc}). "
                      f"Likely a OneDrive cloud-only file, open it once to download it.")
            try:
                dest.unlink(missing_ok=True)
            except OSError:
                pass  # partial copy left behind; harmless, overwritten next run
            continue
        except Exception as exc:
            if verbose:
                print(f"  skip {name}: {type(exc).__name__} on {src.name}: {exc}")
            continue

        found.append(
            LoadedCV(
                name=name,
                file_path=str(dest),
                full_text=full_text,
                bullets=bullets,
                lean_tags=tags,
            )
        )
    return found


def load_master_profile(cv_root: Path) -> str:
    """The superset of honestly-claimable skills. Copied into data/ for grounding."""
    if config.MASTER_PROFILE.exists():
        return config.MASTER_PROFILE.read_text(encoding="utf-8")
    for candidate in cv_root.rglob("MASTER_SKILLS_PROFILE.md"):
        text = candidate.read_text(encoding="utf-8")
        config.MASTER_PROFILE.write_text(text, encoding="utf-8")
        return text
    return ""
